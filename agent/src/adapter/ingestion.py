"""Turning an uploaded or pasted document into text, chunks and figures.

ADAPTER, not core: it imports the PDF and DOCX libraries, so it cannot live
beside the pure policy under ADR-002. What it does NOT do is decide anything -
chunking is `ambassador.knowledge.chunk_text`, figure extraction is
`ambassador.figures.extract_figures`, and both are imported rather than
reimplemented. This module is the bytes-to-text seam and nothing else.

Two rules from docs/10- step 2 shape it. A scanned PDF ends as
`failed/no_extractable_text` rather than an empty document, because an empty
published document is a silence nobody notices while a failure says OCR is
deferred. And original bytes are discarded at the end of the request: the
extracted text is the source of record, and re-parsing means re-uploading.

Nothing here approves a figure. Every occurrence is recorded unapproved and
every chunk is written without a scope, which the repository enforces by taking
no scope argument at all - approval and scope are an admin's, through toby's
review routes.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Final, Literal

import yaml

from ambassador.figures import extract_figures
from ambassador.knowledge import chunk_text, load_limits

_DATA_DIR = Path(__file__).resolve().parents[3] / "data"

SourceType = Literal["pdf", "docx", "txt", "paste"]

ParseErrorCode = Literal[
    "unsupported_type",
    "invalid_encoding",
    "limit_exceeded",
    "no_extractable_text",
    "malformed",
]

_EXTENSIONS: Final[dict[str, SourceType]] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
}

# A sentence, for the figure's context. Deliberately simple: the figure's own
# surface anchors it, so this only has to find the boundaries around it.
_SENTENCE_END = re.compile(r"(?<=[.!?])\s+")

_PARAGRAPH_BREAK: Final = "\n\n"


class ParseFailed(Exception):
    """A parse that failed for a reason worth showing an admin.

    Carries the enum from docs/02-'s `parse_error_code` rather than a message,
    so the route answers with a code the web tier already renders advice for
    (#113) instead of a string it would have to pattern-match.
    """

    def __init__(self, code: ParseErrorCode) -> None:
        super().__init__(code)
        self.code = code


@dataclass(frozen=True)
class ParsedDocument:
    source_type: SourceType
    text: str
    original_filename: str | None
    mime_type: str
    source_bytes: int
    source_sha256: str


@dataclass(frozen=True)
class ExtractedFigure:
    """One occurrence, with the sentence that gives it meaning.

    `active_approval_id` is always None here and is not a parameter: a parser
    that could approve a figure is a parser that eventually does.
    """

    value: str
    kind: str
    currency: str | None
    unit: str | None
    surface: str
    source_sentence: str
    page: int | None = None
    active_approval_id: None = None


def max_source_bytes() -> int:
    """The upload cap, from `data/knowledge.yaml`.

    One number for two tiers: the web route refuses above it twice (#113) and
    this refuses it again, which is the refusal that counts.
    """
    loaded = yaml.safe_load((_DATA_DIR / "knowledge.yaml").read_text(encoding="utf-8"))
    return int(loaded["max_source_bytes"])


def parse_document(
    raw: bytes | None,
    filename: str | None,
    *,
    pasted: str | None = None,
) -> ParsedDocument:
    """Bytes or pasted text in, text out - or `ParseFailed` with a code."""
    if pasted is not None:
        text = pasted.strip()
        if not text:
            raise ParseFailed("no_extractable_text")
        encoded = text.encode("utf-8")
        return ParsedDocument(
            source_type="paste",
            text=text,
            original_filename=None,
            mime_type="text/plain",
            source_bytes=len(encoded),
            source_sha256=hashlib.sha256(encoded).hexdigest(),
        )

    if raw is None or filename is None:
        raise ParseFailed("unsupported_type")
    if len(raw) > max_source_bytes():
        raise ParseFailed("limit_exceeded")

    suffix = Path(filename).suffix.lower()
    source_type = _EXTENSIONS.get(suffix)
    if source_type is None:
        # xlsx, doc, images and URLs are deferred (docs/06-), and a parser that
        # guessed from the bytes would ingest one of them by accident.
        raise ParseFailed("unsupported_type")

    text = _extract(source_type, raw).strip()
    if not text:
        # The scan case. It parsed; there is simply nothing in it.
        raise ParseFailed("no_extractable_text")

    return ParsedDocument(
        source_type=source_type,
        text=text,
        original_filename=filename,
        mime_type=_MIME[source_type],
        source_bytes=len(raw),
        source_sha256=hashlib.sha256(raw).hexdigest(),
    )


_MIME: Final[dict[SourceType, str]] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "paste": "text/plain",
}


def _extract(source_type: SourceType, raw: bytes) -> str:
    if source_type == "txt":
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            raise ParseFailed("invalid_encoding") from None
    if source_type == "pdf":
        return _extract_pdf(raw)
    return _extract_docx(raw)


def _extract_pdf(raw: bytes) -> str:
    """Page order preserved, pages separated as paragraphs.

    Page numbers are kept by the separation rather than by an annotation: the
    chunker works on paragraphs, so a page break that is also a paragraph break
    keeps a chunk from spanning two pages silently.
    """
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except (PdfReadError, ValueError, KeyError, OSError):
        raise ParseFailed("malformed") from None
    return _PARAGRAPH_BREAK.join(page for page in pages if page)


def _extract_docx(raw: bytes) -> str:
    """Paragraphs and table cells in document order (docs/10- step 2)."""
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(raw))
    except (PackageNotFoundError, KeyError, ValueError, OSError):
        raise ParseFailed("malformed") from None

    parts = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _PARAGRAPH_BREAK.join(part for part in parts if part)


def figures_in(text: str) -> list[ExtractedFigure]:
    """Every currency amount and count in the text, as occurrences.

    NOT de-duplicated. The same value written twice is two rows, because
    approval is per occurrence and one approval standing for both would approve
    a sentence nobody read. The extractor is the existing deterministic one, so
    the figures an admin reviews are the same figures the guardrail knows.
    """
    found: list[ExtractedFigure] = []
    for match in extract_figures(text):
        figure = match.figure
        found.append(
            ExtractedFigure(
                value=str(figure.value),
                kind=figure.kind,
                currency=getattr(figure, "currency", None),
                unit=getattr(figure, "unit", None),
                surface=figure.surface,
                source_sentence=_sentence_around(text, match.start),
            )
        )
    return found


def _sentence_around(text: str, position: int) -> str:
    """The sentence a figure sits in, for review.

    docs/10-: approving a value without its sentence and page is not review. So
    this is not decoration - it is the thing being reviewed.
    """
    start = 0
    for boundary in _SENTENCE_END.finditer(text):
        if boundary.end() > position:
            break
        start = boundary.end()
    end = len(text)
    for boundary in _SENTENCE_END.finditer(text, position):
        end = boundary.start()
        break
    return text[start:end].strip()


async def store_document(
    repository: Any,
    *,
    title: str,
    parsed: ParsedDocument,
) -> dict[str, Any]:
    """Write the document, its chunks and every figure, through the repository.

    No SQL here (ADR-021, toby's convention): the repository owns the
    statements, and `add_chunk` deliberately takes no scope argument, so this
    cannot write a chunk that is anything but the default `admin_only`.
    """
    document_id = await repository.add_document(
        revision=1,
        title=title,
        source_type=parsed.source_type,
        original_filename=parsed.original_filename,
        mime_type=parsed.mime_type,
        source_bytes=parsed.source_bytes,
        source_sha256=parsed.source_sha256,
        extracted_text=parsed.text,
    )

    chunks = chunk_text(parsed.text, load_limits())
    written_chunks = 0
    written_figures = 0
    for chunk in chunks:
        chunk_id = await repository.add_chunk(
            document_id=document_id,
            document_revision=1,
            ordinal=chunk.ordinal,
            heading=chunk.heading,
            body=chunk.body,
            content_sha256=hashlib.sha256(chunk.body.encode("utf-8")).hexdigest(),
        )
        written_chunks += 1
        # Figures are found per CHUNK so each one belongs to the chunk whose
        # scope decides whether it could ever be spoken.
        for figure in figures_in(chunk.body):
            await repository.add_figure(
                document_id=document_id,
                document_revision=1,
                chunk_id=chunk_id,
                value=figure.value,
                kind=figure.kind,
                currency=figure.currency,
                unit=figure.unit,
                surface=figure.surface,
                source_sentence=figure.source_sentence,
                page=figure.page,
            )
            written_figures += 1

    return {
        "id": str(document_id),
        "revision": 1,
        "status": "draft",
        "chunks": written_chunks,
        "figures": written_figures,
    }
